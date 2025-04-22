from flask import Flask, request, jsonify, render_template,session
from flask import Flask, send_file, jsonify
import sqlite3
import os
from werkzeug.utils import secure_filename
from docx import Document
import datetime
import shutil
import ollama
import time  
import logging  
import fitz
import json
from docx import Document
from docx.shared import Pt 
import pandas as pd
import re
# import numpy as np
from sklearn.preprocessing import LabelEncoder, StandardScaler
from sklearn.neighbors import NearestNeighbors
from datetime import datetime
from io import BytesIO


# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(message)s')
app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'              
EXPORT_FOLDER = 'exports'  
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(EXPORT_FOLDER, exist_ok=True)


def calculate_maintenance(condition,num_children, nok_details, basic_info):
    TOTAL_ALLOWANCE_CAP = 33  # Max allowance for wife + children
    CHILD_ALLOWANCE_RATE = 5.5  # Default fixed allowance per child when wife gets MA

    # Initialize allowances
    allowance = {"wife": 0, "children": []}
    condition_percentages = {}

    print(f"[DEBUG] Found {num_children} children in calculate_maintenance().")

    # **Helper function to extract max percentage from a range**
    def parse_percentage_range(percentage_range):
        matches = re.findall(r'\d+', percentage_range)
        return float(matches[-1]) if matches else 0

    ### Step 2: Check for Three Conditions Granting 0% MA to Wife ###
    if condition.get("adultery_by_lady") or condition.get("extramarital_by_lady") or condition.get("desertion_without_valid_reason"):
        allowance["wife"] = "Eligible: 10%"
        condition_percentages["No Maintenance - Adultery, Extramarital & Desertion by Lady"] = "10%"

        # **Apply New Rules for Children's MA when Wife Gets 0%**
        allowance["children"] = assign_children_allowance(num_children, nok_details, basic_info)
        return allowance, condition_percentages  # Exit early

    ### Step 3: Calculate Wife’s Allowance Based on Conditions ###
    wife_allowances = [] # returing
    if condition.get("adultery_by_individual"):
        wife_allowances.append("15% - 22%")
        condition_percentages["Adultery by individual"] = "15% - 22%"

    if condition.get("desertion_with_valid_reason"):
        wife_allowances.append("15% - 22%")
        condition_percentages["Desertion with valid reason"] = "15% - 22%"

    # educational qual and earnigs
    if condition.get("lady_is_matriculate/non_matriculate") or condition.get('lady_intermediate'):
        if condition.get("lady_earnings_salaried") or condition.get("lady_earnings_self_employed"):
            if condition.get("income_bracket") == ">8000":
                wife_allowances.append("0%")
                condition_percentages["Lady is Matriculate/Non-mMaytriculate/Itermedite and salaried/self-employed and earns more than ₹8000, no MA"] = "0%"
            else:
                wife_allowances.append("08% - 10%")
                condition_percentages["Lady is Matriculate/Non-mMaytriculate/Itermedite and salaried/self-employed and earns less than ₹8000"] = "08% - 10"

    if condition.get("lady_is_matriculate/non_matriculate") or condition.get('lady_intermediate'):
        if condition.get("lady_earnings_not_working"):
                wife_allowances.append("15% - 22%")
                condition_percentages["Lady is Matriculate/ Non Matriculate and not working"] = "15% - 22%"
        elif condition.get("lady_earnings_prior_work"):
            wife_allowances.append("8% - 10%")
            condition_percentages["Lady is matriculate/Nonmatriculate/ Intermediate and has prior work experience"] = "08% -10"

        if condition.get("lady_graduate") or condition.get('lady_post_graduate'):
            if condition.get("lady_earnings_salaried") or condition.get("lady_earnings_self_employed"):
                if condition.get("income_bracket") == ">8000":
                    wife_allowances.append("0%")
                condition_percentages["Lady is Graduate/Post Graduate and salaried/self-employed and earns more than ₹8000, no MA"] = "0%"
            else:
                wife_allowances.append("08% - 10%")
                condition_percentages["Lady is Graduate/Post Graduate and salaried/self-employed and earns less than ₹8000"] = "08% - 10"

    if condition.get("lady_graduate") or condition.get('lady_post_graduate'):
        if condition.get("lady_earnings_not_working"):
                wife_allowances.append("08% - 10%")
                condition_percentages["Lady is Graduate/Post Graduate and not working"] = "08% - 10%"
        elif condition.get("lady_earnings_prior_work"):
            wife_allowances.append("8% - 10%")
            condition_percentages["Lady is Graduate/Post Graduate  and has prior work experience"] = "08% - 10%"
 
    #refusal of JPC
    if condition.get("lady_refused_jpc"):
        wife_allowances.append("08% - 10%")
        condition_percentages["Lady refused JPC"] = "08% - 10%"

    elif condition.get("individual_refused_jpc"):
        wife_allowances.append("15% - 22%")
        condition_percentages["Individual refused JPC"] = "15% - 22%"

    #fin liability
    if condition.get("liability_no"):
        wife_allowances.append("15% - 22%")
        condition_percentages["No fin Liability on the indl"] = "15% - 22%"

    elif condition.get("liability_personal_car"):
        wife_allowances.append("15% - 22%")
        condition_percentages["Car/ Personal loan"] = "15% - 22%"

    elif condition.get("liability_home_loan"):
        wife_allowances.append("08% - 10%")
        condition_percentages["Home/ DependentParents"] = "08% - 10%"

# Property with lady

    if condition.get("property_movable") or condition.get('property_immovable'):
        if condition.get("property_income_above_10000"):
                wife_allowances.append("0%")
                condition_percentages["Lady's income from property is >10000"] = "0%"
        elif condition.get("property_income_below_10000"):
            wife_allowances.append("8% - 10%")
            condition_percentages["Lady's income from property is <10000"] = "08% - 10%"

    elif condition.get("property_none"):
        wife_allowances.append("8% - 10%")
        condition_percentages["Lady's income from property is nil"] = "08% - 10%"

#Extn of MA

    if condition.get("maint_allce_filed"):
        if condition.get("court_status_subjudice"):
                wife_allowances.append("15% - 22%")
                condition_percentages["Lady's court status case is Subjudice"] = "15% - 22%"
        elif condition.get("court_status_granted"):
            wife_allowances.append("0%")
            condition_percentages["Court has already gtd allce"] = "0%"

    elif condition.get("maint_allce_not_filed"):
        wife_allowances.append("8% - 10%")
        condition_percentages["Lady gtd maint foe 3 years from dt of complaint"] = "08% - 10%"


#Custody
    if condition.get("custody_with_lady"):
        wife_allowances.append("8% - 10%")
        condition_percentages["custody of children is with lady"] = "08% - 10%"

    elif condition.get("custody_with_individual"):
        wife_allowances.append("8% - 10%")
        condition_percentages["custody of children is with lady"] = "08% - 10%"

    # **Select Highest Valid Allowance**
    max_wife_allowance = max(wife_allowances, key=parse_percentage_range, default="0%")
    remaining_allowance = TOTAL_ALLOWANCE_CAP - (num_children * CHILD_ALLOWANCE_RATE)

    # **Ensure It Doesn't Exceed the Remaining Allowance**
    max_wife_allowance_float = parse_percentage_range(max_wife_allowance)
    if max_wife_allowance_float <= remaining_allowance:
        allowance["wife"] = f"Eligible: {max_wife_allowance}"
    else:
        allowance["wife"] = f"Eligible: {remaining_allowance}%"

    ### Step 4: Calculate Children's Maintenance Allowance ###
    if parse_percentage_range(allowance["wife"]) > 0:
        allowance["children"] = [f"Eligible: {CHILD_ALLOWANCE_RATE}%"] * num_children

    else:
        allowance["children"] = assign_children_allowance(num_children, nok_details, basic_info)

    return allowance, condition_percentages


def assign_children_allowance(num_children, nok_details, basic_info):
    """ Determines children's maintenance allowance based on dependency, age, education, and HRA city classification. """
    children_allowance = []

    # **Step 1: Filter Only Sons & Daughters**
    print(f"[DEBUG] Processing {num_children} children in assign_children_allowance().")

    for child in num_children:
        age = int(child.get("age", 0))
        education = child.get("education", "").strip().lower()        
        dependent_on_mother = child.get("dependent_on_mother", "").strip().lower() == "yes"
        hra_city = basic_info.get("hra_city", "").strip().upper()

        if dependent_on_mother and (15 <= age <= 25) and ("intermediate" in education or "graduation" in education) and hra_city == "X":
            children_allowance.append("Eligible: 10% - 25%")
        elif dependent_on_mother and (15 <= age <= 25) and ("intermediate" in education or "graduation" in education) and hra_city in ["Y", "Z"]:
            children_allowance.append("Eligible: 6% - 10%")
        elif dependent_on_mother and age < 15:
            children_allowance.append("Eligible: 8% - 10%")
        elif not dependent_on_mother and hra_city == "X":
            children_allowance.append("Eligible: 10% - 25%")
        elif not dependent_on_mother and hra_city in ["Y", "Z"]:
            children_allowance.append("Eligible: 6% - 10%")
        else:
            children_allowance.append("Eligible: 5.5%")

    return children_allowance



# Route to render the HTML page
@app.route('/')
def index():
    return render_template('index.html')



def create_tables():
    """Ensure all tables exist before inserting data."""
    conn = sqlite3.connect('database.db')
    c = conn.cursor()

    # Table: basic_info
    c.execute('''
        CREATE TABLE IF NOT EXISTS basic_info (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            army_number TEXT UNIQUE NOT NULL, 
            unit TEXT, 
            rank TEXT, 
            name TEXT, 
            pcda_acct TEXT,  
            date_of_enrolment TEXT, 
            date_of_present_rank TEXT, 
            date_of_retirement TEXT, 
            date_of_marriage TEXT, 
            hra TEXT, 
            city_name TEXT, 
            spr TEXT
        )
    ''')

    # Table: legal_cases
    c.execute('''
        CREATE TABLE IF NOT EXISTS legal_cases (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            army_number TEXT UNIQUE NOT NULL, 
            date_of_complaint TEXT, 
            complaint_details TEXT, 
            court_case TEXT, 
            case_description TEXT,
            FOREIGN KEY (army_number) REFERENCES basic_info(army_number)
        )
    ''')

    # Table: lady_details
    c.execute('''
        CREATE TABLE IF NOT EXISTS lady_details (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            army_number TEXT UNIQUE NOT NULL, 
            lady_name TEXT, 
            lady_address TEXT, 
            lady_banker TEXT, 
            lady_acct_number TEXT, 
            lady_ifsc TEXT, 
            date_of_affidavit TEXT,
            FOREIGN KEY (army_number) REFERENCES basic_info(army_number)
        )
    ''')

    # Table: uploads (stores affidavit and other files)
    c.execute('''
        CREATE TABLE IF NOT EXISTS uploads (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            army_number TEXT UNIQUE NOT NULL, 
            upload_affidavit TEXT,
            FOREIGN KEY (army_number) REFERENCES basic_info(army_number)
        )
    ''')

    conn.commit()
    conn.close()



# Call the function to create tables before inserting data
create_tables()



@app.route('/submit_basic_info_officer', methods=['POST'])
def submit_basic_info():
    try:
        print("[DEBUG] Received request to submit basic info")
        
        army_number = request.form['army-number']
        unit = request.form['unit']
        rank = request.form['rank']
        name = request.form['name']
        pcda_acct = request.form.get('pcda-acct') or None
        date_of_enrolment = request.form.get('date-of-enrolment') or None

        date_of_present_rank = request.form['date-of-present-rank']
        date_of_retirement = request.form['date-of-retirement']
        date_of_marriage = request.form['date-of-marriage']
        hra = request.form['hra']
        city_name = request.form['city-name']
        spr = request.form['spr']

        # Legal Case Data
        date_of_complaint = request.form['date-of-complaint']
        complaint_details = request.form.get('complaint-details', '')
        court_case = request.form['court-case']
        case_description = request.form.get('case-description', '')

        # Lady Details
        lady_name = request.form['lady-name']
        lady_address = request.form['lady-address']
        lady_banker = request.form['lady-banker']
        lady_acct_number = request.form['lady-acct-number']
        lady_ifsc = request.form['lady-ifsc']
        date_of_affidavit = request.form['date-of-affidavit']

        # File Uploads
        upload_affidavit = request.files['upload-affidavit'].filename 

        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        print("[DEBUG] Database connected")

        # Insert into basic_info
        c.execute('''
            INSERT OR REPLACE INTO basic_info (army_number, unit, rank, name, pcda_acct, 
                                               date_of_enrolment, date_of_present_rank, 
                                               date_of_retirement, date_of_marriage, hra, 
                                               city_name, spr)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (army_number, unit, rank, name, pcda_acct, date_of_enrolment, 
              date_of_present_rank, date_of_retirement, date_of_marriage, hra, 
              city_name, spr))

        # Insert into legal_cases
        c.execute('''
            INSERT OR REPLACE INTO legal_cases (army_number, date_of_complaint, 
                                                complaint_details, court_case, case_description)
            VALUES (?, ?, ?, ?, ?)
        ''', (army_number, date_of_complaint, complaint_details, court_case, case_description))

        # Insert into lady_details
        c.execute('''
            INSERT OR REPLACE INTO lady_details (army_number, lady_name, lady_address, 
                                                 lady_banker, lady_acct_number, lady_ifsc, 
                                                 date_of_affidavit)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (army_number, lady_name, lady_address, lady_banker, 
              lady_acct_number, lady_ifsc, date_of_affidavit))

        # Insert into uploads
        c.execute('''
            INSERT OR REPLACE INTO uploads (army_number, upload_affidavit)
            VALUES (?, ?)
        ''', (army_number, upload_affidavit))

        conn.commit()
        print("[DEBUG] Data committed to database")
        conn.close()
        print("[DEBUG] Database connection closed")

        return jsonify({'success': True, 'message': 'Data saved successfully!'})

    except Exception as e:
        print(f"[ERROR] Exception occurred: {e}")
        return jsonify({'success': False, 'message': str(e)})

@app.route('/submit_basic_info_jco', methods=['POST'])
def submit_basic_info_jco():
    try:
        print("[DEBUG] Received request to submit basic info")
        
        army_number = request.form['army-number']
        unit = request.form['unit']
        rank = request.form['rank']
        name = request.form['name']
        pcda_acct = request.form.get('pcda-acct') or None
        date_of_enrolment = request.form.get('date-of-enrolment') or None

        date_of_present_rank = request.form['date-of-present-rank']
        date_of_retirement = request.form['date-of-retirement']
        date_of_marriage = request.form['date-of-marriage']
        hra = request.form['hra']
        city_name = request.form['city-name']
        spr = request.form['spr']

        # Legal Case Data
        date_of_complaint = request.form['date-of-complaint']
        complaint_details = request.form.get('complaint-details', '')
        court_case = request.form['court-case']
        case_description = request.form.get('case-description', '')

        # Lady Details
        lady_name = request.form['lady-name']
        lady_address = request.form['lady-address']
        lady_banker = request.form['lady-banker']
        lady_acct_number = request.form['lady-acct-number']
        lady_ifsc = request.form['lady-ifsc']
        date_of_affidavit = request.form['date-of-affidavit']

        # File Uploads
        upload_affidavit = request.files['upload-affidavit'].filename 

        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        print("[DEBUG] Database connected")

        # Insert into basic_info
        c.execute('''
            INSERT OR REPLACE INTO basic_info (army_number, unit, rank, name, pcda_acct, 
                                               date_of_enrolment, date_of_present_rank, 
                                               date_of_retirement, date_of_marriage, hra, 
                                               city_name, spr)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (army_number, unit, rank, name, pcda_acct, date_of_enrolment, 
              date_of_present_rank, date_of_retirement, date_of_marriage, hra, 
              city_name, spr))

        # Insert into legal_cases
        c.execute('''
            INSERT OR REPLACE INTO legal_cases (army_number, date_of_complaint, 
                                                complaint_details, court_case, case_description)
            VALUES (?, ?, ?, ?, ?)
        ''', (army_number, date_of_complaint, complaint_details, court_case, case_description))

        # Insert into lady_details
        c.execute('''
            INSERT OR REPLACE INTO lady_details (army_number, lady_name, lady_address, 
                                                 lady_banker, lady_acct_number, lady_ifsc, 
                                                 date_of_affidavit)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (army_number, lady_name, lady_address, lady_banker, 
              lady_acct_number, lady_ifsc, date_of_affidavit))

        # Insert into uploads
        c.execute('''
            INSERT OR REPLACE INTO uploads (army_number, upload_affidavit)
            VALUES (?, ?)
        ''', (army_number, upload_affidavit))

        conn.commit()
        print("[DEBUG] Data committed to database")
        conn.close()
        print("[DEBUG] Database connection closed")

        return jsonify({'success': True, 'message': 'Data saved successfully!'})

    except Exception as e:
        print(f"[ERROR] Exception occurred: {e}")
        return jsonify({'success': False, 'message': str(e)})


@app.route('/submit_nok_details', methods=['POST'])
def submit_nok_details():
    try:
        print("\n[DEBUG] Received request to submit NOK details.\n")

        # First, try to fetch army_number from request
        army_number = request.json.get('army_number')

        # If army_number is not provided, fetch the latest army_number from basic_info
        if not army_number:
            conn = sqlite3.connect('database.db')
            c = conn.cursor()
            c.execute("SELECT army_number FROM basic_info ORDER BY id DESC LIMIT 1")
            result = c.fetchone()
            if result:
                army_number = result[0]
            conn.close()
        
        print(f"[DEBUG] Army Number Fetched: {army_number}")

        # Fetch NOK details from request
        nok_data = request.json.get('nokData', [])

        if not army_number or not nok_data:
            print("[DEBUG] No valid army_number or NOK data received.")
            return jsonify({'success': False, 'message': 'Army number or NOK data missing.'})

        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        print("[DEBUG] Database connected.")

        # Ensure the table exists
        c.execute('''
            CREATE TABLE IF NOT EXISTS nok_details (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                army_number TEXT NOT NULL,  
                name TEXT, 
                relation TEXT, 
                date_of_birth TEXT, 
                age INTEGER, 
                education TEXT DEFAULT 'N/A',  
                part_ii_child TEXT DEFAULT 'N/A',  
                dependent_on_mother TEXT DEFAULT 'N/A',  
                part_ii_marriage TEXT DEFAULT 'N/A',  
                remarks TEXT,
                FOREIGN KEY (army_number) REFERENCES basic_info(army_number) ON DELETE CASCADE
            )
        ''')
        print("[DEBUG] NOK Details table checked/created.")

        saved_data = []  # List to store inserted data for saving

        # Insert each NOK entry
        for nok in nok_data:
            print(f"\n[DEBUG] Processing NOK entry: {nok}\n")

            c.execute('''
                INSERT INTO nok_details (army_number, name, relation, date_of_birth, age, 
                                         education, part_ii_child, dependent_on_mother, 
                                         part_ii_marriage, remarks)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (
                army_number, nok['name'], nok['relation'], nok['dob'], nok['age'], 
                nok.get('education', "N/A"),  
                nok.get('part_ii_child', "N/A"),  
                nok.get('dependent_on_mother', "N/A"),  
                nok.get('part_ii_marriage', "N/A"),  
                nok['remarks']
            ))
            print("[DEBUG] NOK entry inserted into database.")

            # Save the inserted data into a list
            saved_data.append(nok)

        session['nok'] = saved_data
        conn.commit()
        print("[DEBUG] Database commit successful.")
        conn.close()
        print("[DEBUG] Database connection closed.")

        # Save the inserted data separately in a JSON file
        with open('saved_nok_details.json', 'a') as f:
            f.write(json.dumps(saved_data, indent=4) + "\n")
        print("[DEBUG] NOK details saved in saved_nok_details.json.")

        return jsonify({'success': True, 'message': 'NOK details saved successfully!'})

    except Exception as e:
        print(f"[ERROR] Exception occurred: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500





@app.route('/submit_facilities', methods=['POST'])
def submit_facilities():
    try:
        print("\n[DEBUG] Received request to submit facilities availed.\n")

        # Fetch army_number from request
        army_number = request.form.get('army-number')

        # If army_number is not provided, fetch the latest army_number from basic_info
        if not army_number:
            conn = sqlite3.connect('database.db')
            c = conn.cursor()
            c.execute("SELECT army_number FROM basic_info ORDER BY id DESC LIMIT 1")
            result = c.fetchone()
            if result:
                army_number = result[0]
            conn.close()

        print(f"[DEBUG] Army Number Fetched: {army_number}")

        # Validate army_number
        if not army_number:
            print("[ERROR] Army number is missing!")
            return jsonify({'success': False, 'message': 'Army number is required.'}), 400

        # Fetch the basic facilities data
        data = {
            'canteen_card': request.form.get('canteen-card', 'no'),
            'mec': request.form.get('mec', 'no'),
            'dependent_card': request.form.get('dependent-card', 'no'),
            'joint_acct': request.form.get('joint-acct', 'no'),
            'sf_accommodation': request.form.get('sf-accommodation', 'no'),
        }

        print(f"[DEBUG] Parsed Basic Facilities Data: {data}")

        # If Joint Account is Yes, fetch the additional details
        if data['joint_acct'] == 'yes':
            joint_acct_details = {
                'joint_acct_no': request.form.get('joint-acct-no', ''),
                'joint_banker': request.form.get('joint-banker', ''),
                'joint_ifsc': request.form.get('joint-ifsc', ''),
            }
            print("[DEBUG] Joint Account is YES, Capturing Additional Details.")
        else:
            joint_acct_details = {
                'joint_acct_no': None,
                'joint_banker': None,
                'joint_ifsc': None,
            }
            print("[DEBUG] Joint Account is NO, Skipping Additional Details.")

        # Combine all data
        full_data = {'army_number': army_number, **data, **joint_acct_details}
        print(f"[DEBUG] Final Data to be Inserted: {full_data}")

        # Save the data into the database
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        print("[DEBUG] Database Connected.")

        # Ensure the facilities table exists
        c.execute('''
            CREATE TABLE IF NOT EXISTS facilities (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                army_number TEXT NOT NULL,  
                canteen_card TEXT, mec TEXT, dependent_card TEXT,
                joint_acct TEXT, joint_acct_no TEXT, joint_banker TEXT,
                joint_ifsc TEXT, sf_accommodation TEXT,
                FOREIGN KEY (army_number) REFERENCES basic_info(army_number) ON DELETE CASCADE
            )
        ''')
        print("[DEBUG] Facilities Table Checked/Created.")

        # Insert or update facilities data
        c.execute('''
            INSERT OR REPLACE INTO facilities (
                army_number, canteen_card, mec, dependent_card, joint_acct,
                joint_acct_no, joint_banker, joint_ifsc, sf_accommodation
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            full_data['army_number'], full_data['canteen_card'], full_data['mec'],
            full_data['dependent_card'], full_data['joint_acct'], full_data['joint_acct_no'],
            full_data['joint_banker'], full_data['joint_ifsc'], full_data['sf_accommodation']
        ))
        print("[DEBUG] Facilities Data Inserted into Database.")

        conn.commit()
        print("[DEBUG] Database Commit Successful.")
        conn.close()
        print("[DEBUG] Database Connection Closed.")

        # Save inserted data separately in a JSON file for future reference
        with open('saved_facilities.json', 'a') as f:
            f.write(json.dumps(full_data, indent=4) + "\n")
        print("[DEBUG] Facilities details saved in saved_facilities.json.")

        return jsonify({'success': True, 'message': 'Facilities availed data saved successfully!'})

    except Exception as e:
        print(f"[ERROR] Exception occurred: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500



# Route to fetch data for a given Army Number
@app.route('/fetch_details', methods=['GET'])
def fetch_details():
    army_number = request.args.get('army_number')

    if not army_number:
        return jsonify({'success': False, 'message': 'Army Number is required.'})

    conn = sqlite3.connect('database.db')
    c = conn.cursor()
    
    # Fetch Individual Details
    c.execute('SELECT * FROM basic_info WHERE army_number = ?', (army_number,))
    row = c.fetchone()

    if not row:
        conn.close()
        return jsonify({'success': False, 'message': 'No details found for the given Army Number.'})

    # Define field names
    keys = [
        'id', 'army_number', 'unit', 'rank', 'name', 'date_of_enrolment',
        'date_of_present_rank', 'date_of_retirement', 'date_of_complaint',
        'date_of_marriage', 'hra', 'city_name', 'spr', 'court_case', 'case_description',
        'lady_address', 'lady_banker', 'lady_acct_number', 'lady_ifsc'
    ]
    
    # Convert row to dictionary
    data = dict(zip(keys, row))

    conn.close()
    return jsonify({'success': True, **data})



# Ensure upload directory exists
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)



def save_file(file):
    """Save uploaded file and return its path."""
    if file and file.filename:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)
        return file.filename  # Store only the filename
    return None




@app.route('/submit_documents', methods=['POST'])
def submit_documents():
    try:
        print("\n[DEBUG] Received request to submit documents.\n")

        # Fetch army_number from request
        army_number = request.form.get('army-number')

        # If army_number is not provided, fetch the latest army_number from basic_info
        if not army_number:
            conn = sqlite3.connect('database.db')
            c = conn.cursor()
            c.execute("SELECT army_number FROM basic_info ORDER BY id DESC LIMIT 1")
            result = c.fetchone()
            if result:
                army_number = result[0]
            conn.close()

        print(f"[DEBUG] Army Number Fetched: {army_number}")

        # Validate army_number
        if not army_number:
            print("[ERROR] Army number is missing!")
            return jsonify({'success': False, 'message': 'Army number is required.'}), 400

        # File uploads
        outcome_counselling_file = request.files.get('outcome-counselling')
        effort_reconciliation_file = request.files.get('effort-reconciliation')
        brief_case_file = request.files.get('brief-case')
        pay_slip_file = request.files.get('pay-slip')

        reply_pre_scn_file = request.files.get('reply-pre-scn')  # New field
        reply_scn_file = request.files.get('reply-scn')  # New field

        # Manual entries (date fields)
        date_pre_scn = request.form.get('date-pre-scn', '')
        date_reply_pre_scn = request.form.get('date-reply-pre-scn', '')
        date_scn = request.form.get('date-scn', '')
        date_reply_scn = request.form.get('date-reply-scn', '')

        print(f"[DEBUG] Parsed Date Fields: Pre SCN: {date_pre_scn}, Reply Pre SCN: {date_reply_pre_scn}, SCN: {date_scn}, Reply SCN: {date_reply_scn}")

        # Save uploaded files
        outcome_counselling_path = save_file(outcome_counselling_file)
        effort_reconciliation_path = save_file(effort_reconciliation_file)
        brief_case_path = save_file(brief_case_file)
        pay_slip_path = save_file(pay_slip_file)
        reply_pre_scn_path = save_file(reply_pre_scn_file)  # New field
        reply_scn_path = save_file(reply_scn_file)  # New field

        print(f"[DEBUG] File paths stored: Outcome Counselling: {outcome_counselling_path}, Effort Reconciliation: {effort_reconciliation_path}, Brief Case: {brief_case_path}, Pay Slip: {pay_slip_path}, Reply Pre SCN: {reply_pre_scn_path}, Reply SCN: {reply_scn_path}")

        # Save details to database
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        print("[DEBUG] Database Connected.")

        # Ensure documents table exists
        c.execute('''
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                army_number TEXT NOT NULL,  
                outcome_counselling TEXT, effort_reconciliation TEXT, brief_case TEXT,
                pay_slip TEXT, 
                date_pre_scn TEXT, date_reply_pre_scn TEXT, reply_pre_scn TEXT,
                date_scn TEXT, date_reply_scn TEXT, reply_scn TEXT,
                FOREIGN KEY (army_number) REFERENCES basic_info(army_number) ON DELETE CASCADE
            )
        ''')
        print("[DEBUG] Documents Table Checked/Created.")

        # Insert data into the database
        c.execute('''
            INSERT OR REPLACE INTO documents (
                army_number, outcome_counselling, effort_reconciliation, brief_case,
                pay_slip, date_pre_scn, date_reply_pre_scn, reply_pre_scn,
                date_scn, date_reply_scn, reply_scn
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', (
            army_number, outcome_counselling_path, effort_reconciliation_path, brief_case_path,
            pay_slip_path, date_pre_scn, date_reply_pre_scn, reply_pre_scn_path,
            date_scn, date_reply_scn, reply_scn_path
        ))
        print("[DEBUG] Documents Data Inserted into Database.")

        conn.commit()
        print("[DEBUG] Database Commit Successful.")
        conn.close()
        print("[DEBUG] Database Connection Closed.")

        # Save inserted data separately in a JSON file for future reference
        with open('saved_documents.json', 'a') as f:
            f.write(json.dumps({
                "army_number": army_number,
                "outcome_counselling": outcome_counselling_path,
                "effort_reconciliation": effort_reconciliation_path,
                "brief_case": brief_case_path,
                "pay_slip": pay_slip_path,
                "date_pre_scn": date_pre_scn,
                "date_reply_pre_scn": date_reply_pre_scn,
                "reply_pre_scn": reply_pre_scn_path,
                "date_scn": date_scn,
                "date_reply_scn": date_reply_scn,
                "reply_scn": reply_scn_path
            }, indent=4) + "\n")
        print("[DEBUG] Documents details saved in saved_documents.json.")

        return jsonify({'success': True, 'message': 'Documents uploaded and saved successfully!'})

    except Exception as e:
        print(f"[ERROR] Exception occurred: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500



app.secret_key = "your_secret_key"



@app.route('/submit_matrix', methods=['POST'])
def submit_matrix():
    try:
        print("\n[DEBUG] Received request to submit matrix data.\n")

        # Ensure request is JSON
        if not request.is_json:
            print("[ERROR] Invalid JSON format received.")
            return jsonify({'success': False, 'message': 'Invalid JSON format! Expected JSON but received something else.'}), 400

        # Parse JSON data correctly
        data = request.get_json()
        if not data:
            print("[ERROR] No matrix data received.")
            return jsonify({'success': False, 'message': 'No matrix data received!'}), 400

        army_number = data.get('army_number')

        # If army_number is not provided, fetch the latest army_number from basic_info
        if not army_number:
            conn = sqlite3.connect('database.db')
            c = conn.cursor()
            c.execute("SELECT army_number FROM basic_info ORDER BY id DESC LIMIT 1")
            result = c.fetchone()
            if result:
                army_number = result[0]
            conn.close()

        print(f"[DEBUG] Army Number Fetched: {army_number}")

        # Validate army_number
        if not army_number:
            print("[ERROR] Army number is missing!")
            return jsonify({'success': False, 'message': 'Army number is required!'}), 400

        # Handle multiple selections (Adultery Checkboxes)
        adultery_selections = data.get('adultery', [])
        if isinstance(adultery_selections, str):  # Convert single selection to list
            adultery_selections = [adultery_selections]

        # Define all expected boolean fields
        processed_data = {
            'army_number': army_number,  # Now uses army_number instead of basic_info ID
            'adultery_by_lady': "Adultery by Lady" in adultery_selections,
            'adultery_by_individual': "Adultery by Indl" in adultery_selections,
            'desertion_with_valid_reason': data.get('desertion') == "With Valid Reason",
            'desertion_without_valid_reason': data.get('desertion') == "Without Valid Reason",
            'lady_earnings_salaried': data.get('lady-earnings') == "Salaried",
            'lady_earnings_self_employed': data.get('lady-earnings') == "Entrepreneur",
            'lady_earnings_prior_work': data.get('lady-earnings') == "Prior work experience",
            'lady_earnings_not_working': data.get('lady-earnings') == "Not Working",
            'income_bracket_above_8000': data.get('income-bracket') == ">8000",
            'income_bracket_below_8000': data.get('income-bracket') == "<8000",
            'lady_matriculate': data.get('lady-education') == "Matriculate/Non-Matriculate",
            'lady_intermediate': data.get('lady-education') == "Intermediate",
            'lady_graduate': data.get('lady-education') == "Graduate",
            'lady_post_graduate': data.get('lady-education') == "Post Graduate",
            'property_none': data.get('property') == "None held",
            'property_movable': data.get('property') == "Movable",
            'property_immovable': data.get('property') == "Immovable",
            'property_income_above_10000': data.get('property-income') == ">10000",
            'property_income_below_10000': data.get('property-income') == "<10000",
            'liability_no': data.get('liability') == "No Liability",
            'liability_personal_car': data.get('liability') == "Personal/Car Loans",
            'liability_home_loan': data.get('liability') == "Home Loan/Parents",
            'maint_allce_filed': data.get('maint-allce') == "Filed in Time",
            'maint_allce_not_filed': data.get('maint-allce') == "Not Filed",
            'court_status_subjudice': data.get('court-status') == "Subjudice",
            'court_status_granted': data.get('court-status') == "Court Granted Maintenance",
            'custody_with_lady': data.get('custody') == "With the Lady",
            'custody_with_individual': data.get('custody') == "With the Indl",
            'lady_refused_jpc': data.get('jpc') == "Lady Refused",
            'individual_refused_jpc': data.get('jpc') == "Indl Refused",
        }

        # Store in session
        session['true_conditions'] = processed_data
        print("[DEBUG] Stored True Conditions in Session:\n", processed_data)

        conn = sqlite3.connect('database.db')
        c = conn.cursor()

        # Ensure table exists
        c.execute('''
            CREATE TABLE IF NOT EXISTS case_matrix (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                army_number TEXT NOT NULL UNIQUE,  
                adultery_by_lady BOOLEAN, adultery_by_individual BOOLEAN,
                desertion_with_valid_reason BOOLEAN, desertion_without_valid_reason BOOLEAN,
                lady_earnings_salaried BOOLEAN, lady_earnings_self_employed BOOLEAN,
                lady_earnings_prior_work BOOLEAN, lady_earnings_not_working BOOLEAN,
                income_bracket_above_8000 BOOLEAN, income_bracket_below_8000 BOOLEAN,
                lady_matriculate BOOLEAN, lady_intermediate BOOLEAN, lady_graduate BOOLEAN, lady_post_graduate BOOLEAN,
                property_none BOOLEAN, property_movable BOOLEAN, property_immovable BOOLEAN,
                property_income_above_10000 BOOLEAN, property_income_below_10000 BOOLEAN,
                liability_no BOOLEAN, liability_personal_car BOOLEAN, liability_home_loan BOOLEAN,
                maint_allce_filed BOOLEAN, maint_allce_not_filed BOOLEAN,
                court_status_subjudice BOOLEAN, court_status_granted BOOLEAN,
                custody_with_lady BOOLEAN, custody_with_individual BOOLEAN,
                lady_refused_jpc BOOLEAN, individual_refused_jpc BOOLEAN,
                FOREIGN KEY (army_number) REFERENCES basic_info(army_number) ON DELETE CASCADE
            )
        ''')

        # Insert/Update data
        placeholders = ', '.join('?' * len(processed_data))
        query = f"INSERT OR REPLACE INTO case_matrix ({', '.join(processed_data.keys())}) VALUES ({placeholders})"
        c.execute(query, tuple(processed_data.values()))

        conn.commit()
        conn.close()
        print("[DEBUG] Matrix data saved in database.")

        return jsonify({'success': True, 'message': 'Matrix data saved successfully!', 'data': processed_data})

    except Exception as e:
        print("[ERROR] Matrix Saving Error:", str(e))
        return jsonify({'success': False, 'message': str(e)}), 500




def get_matrix_data():
    return session.get("true_conditions", {})




@app.route('/submit_analysis', methods=['POST'])
def submit_analysis():
    try:
        print("\n[DEBUG] Received request to submit analysis data.\n")

        # Fetch army_number from request
        army_number = request.form.get('army-number')

        # If army_number is not provided, fetch the latest army_number from basic_info
        if not army_number:
            conn = sqlite3.connect('database.db')
            c = conn.cursor()
            c.execute("SELECT army_number FROM basic_info ORDER BY id DESC LIMIT 1")
            result = c.fetchone()
            if result:
                army_number = result[0]
            conn.close()

        print(f"[DEBUG] Army Number Fetched: {army_number}")

        # Validate army_number
        if not army_number:
            print("[ERROR] Army number is missing!")
            return jsonify({'success': False, 'message': 'Army number is required!'}), 400

        # File uploads
        psych_counselor_file = request.files.get('psych-counselor')
        bde_sect_cdr_file = request.files.get('bde-sect-cdr')
        div_goc_file = request.files.get('div-goc')
        corps_goc_file = request.files.get('corps-goc')

        # Save uploaded files
        psych_counselor_path = save_file(psych_counselor_file)
        bde_sect_cdr_path = save_file(bde_sect_cdr_file)
        div_goc_path = save_file(div_goc_file)
        corps_goc_path = save_file(corps_goc_file)

        print(f"[DEBUG] File paths stored: Psych Counselor: {psych_counselor_path}, Bde Sect Cdr: {bde_sect_cdr_path}, Div GOC: {div_goc_path}, Corps GOC: {corps_goc_path}")

        # Save details to database
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        print("[DEBUG] Database Connected.")

        # Ensure analysis table exists
        c.execute('''
            CREATE TABLE IF NOT EXISTS analysis (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                army_number TEXT NOT NULL UNIQUE,  
                psych_counselor TEXT, 
                bde_sect_cdr TEXT, 
                div_goc TEXT, 
                corps_goc TEXT,
                FOREIGN KEY (army_number) REFERENCES basic_info(army_number) ON DELETE CASCADE
            )
        ''')
        print("[DEBUG] Analysis Table Checked/Created.")

        # Insert or update analysis data
        c.execute('''
            INSERT OR REPLACE INTO analysis (
                army_number, psych_counselor, bde_sect_cdr, div_goc, corps_goc
            ) VALUES (?, ?, ?, ?, ?)
        ''', (
            army_number, psych_counselor_path, bde_sect_cdr_path, div_goc_path, corps_goc_path
        ))
        print("[DEBUG] Analysis Data Inserted into Database.")

        conn.commit()
        print("[DEBUG] Database Commit Successful.")
        conn.close()
        print("[DEBUG] Database Connection Closed.")

        # Save inserted data separately in a JSON file for future reference
        with open('saved_analysis.json', 'a') as f:
            f.write(json.dumps({
                "army_number": army_number,
                "psych_counselor": psych_counselor_path,
                "bde_sect_cdr": bde_sect_cdr_path,
                "div_goc": div_goc_path,
                "corps_goc": corps_goc_path
            }, indent=4) + "\n")
        print("[DEBUG] Analysis details saved in saved_analysis.json.")

        return jsonify({'success': True, 'message': 'Analysis data uploaded and saved successfully!'})

    except Exception as e:
        print(f"[ERROR] Exception occurred: {e}")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/submit_additional_inputs', methods=['POST'])
def submit_additional_inputs():
    try:
        print("\n[DEBUG] Received request to submit additional inputs.\n")

        # Fetch army_number from request
        army_number = request.form.get('army-number')

        # If army_number is not provided, fetch the latest army_number from basic_info
        if not army_number:
            conn = sqlite3.connect('database.db')
            c = conn.cursor()
            c.execute("SELECT army_number FROM basic_info ORDER BY id DESC LIMIT 1")
            result = c.fetchone()
            if result:
                army_number = result[0]
            conn.close()

        print(f"[DEBUG] Army Number Fetched: {army_number}")

        # Validate army_number (it should not be empty)
        if not army_number:
            print("[ERROR] Army number is missing!")
            return jsonify({'success': False, 'message': 'Army number is required!'}), 400

        # Manual inputs (assuming "Nil" if no input is received, except for army_number)
        data = {
            'army_number': army_number,
            'source_income_wife': request.form.get('source-income-wife', 'Nil'),
            'employment_details': request.form.get('employment-details', 'Nil'),
            'support_details': request.form.get('support-details', 'Nil'),
            'custody_children': request.form.get('custody-children', 'Nil'),
            'living_standard': request.form.get('living-standard', 'Nil'),
            'health_wife': request.form.get('health-wife', 'Nil'),
            'health_children': request.form.get('health-children', 'Nil'),
            'education_expenses': request.form.get('education-expenses', 'Nil'),
            'shelter_wife': request.form.get('shelter-wife', 'Nil'),
            'shelter_children': request.form.get('shelter-children', 'Nil'),
            'expenses_liabilities_individual': request.form.get('expenses-liabilities-individual', 'Nil'),
            'expenses_liabilities_spouse': request.form.get('expenses-liabilities-spouse', 'Nil'),
            'evaluation_abuse': request.form.get('evaluation-abuse', 'Nil'),
            'chances_reconciliation': request.form.get('chances-reconciliation', 'Nil'),
        }

        # Debugging log
        print("[DEBUG] Processed Additional Inputs:\n", json.dumps(data, indent=4))

        # Save details to the database
        conn = sqlite3.connect('database.db')
        c = conn.cursor()

        # Ensure additional_inputs table exists
        c.execute('''
            CREATE TABLE IF NOT EXISTS additional_inputs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                army_number TEXT NOT NULL UNIQUE,  
                source_income_wife TEXT DEFAULT 'Nil',
                employment_details TEXT DEFAULT 'Nil',
                support_details TEXT DEFAULT 'Nil',
                custody_children TEXT DEFAULT 'Nil',
                living_standard TEXT DEFAULT 'Nil',
                health_wife TEXT DEFAULT 'Nil',
                health_children TEXT DEFAULT 'Nil',
                education_expenses TEXT DEFAULT 'Nil',
                shelter_wife TEXT DEFAULT 'Nil',
                shelter_children TEXT DEFAULT 'Nil',
                expenses_liabilities_individual TEXT DEFAULT 'Nil',
                expenses_liabilities_spouse TEXT DEFAULT 'Nil',
                evaluation_abuse TEXT DEFAULT 'Nil',
                chances_reconciliation TEXT DEFAULT 'Nil',
                FOREIGN KEY (army_number) REFERENCES basic_info(army_number) ON DELETE CASCADE
            )
        ''')
        print("[DEBUG] Additional Inputs Table Checked/Created.")

        # **Fix: Ensure all columns (excluding `id`) match the values provided**
        c.execute('''
            INSERT OR REPLACE INTO additional_inputs (
                army_number, source_income_wife, employment_details,
                support_details, custody_children, living_standard,
                health_wife, health_children, education_expenses,
                shelter_wife, shelter_children,
                expenses_liabilities_individual, expenses_liabilities_spouse,
                evaluation_abuse, chances_reconciliation
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ''', tuple(data.values()))

        print("[DEBUG] Additional Inputs Data Inserted into Database.")

        conn.commit()
        conn.close()
        print("[DEBUG] Database Commit Successful.")
        print("[DEBUG] Database Connection Closed.")

        # Save inserted data separately in a JSON file for future reference
        with open('saved_additional_inputs.json', 'a') as f:
            f.write(json.dumps(data, indent=4) + "\n")
        print("[DEBUG] Additional Inputs details saved in saved_additional_inputs.json.")

        return jsonify({'success': True, 'message': 'Additional inputs saved successfully!'})

    except Exception as e:
        print("[ERROR] Additional Inputs Saving Error:", str(e))
        return jsonify({'success': False, 'message': str(e)}), 500


def save_file(uploaded_file):
    """Save uploaded file and return its path."""
    if uploaded_file and uploaded_file.filename:
        filepath = os.path.join(UPLOAD_FOLDER, uploaded_file.filename)
        uploaded_file.save(filepath)
        print(f"[DEBUG] File saved: {filepath}")  # Debugging log
        return filepath
    return None  # Return None if no file was uploaded



@app.route('/fetch_ma_details/<army_number>', methods=['GET'])
def fetch_ma_details(army_number):
    try:
        conn = sqlite3.connect('database.db')
        c = conn.cursor()

        # Fetch conditions from case_matrix
        c.execute("SELECT * FROM case_matrix WHERE army_number = ?", (army_number,))
        case_data = c.fetchone()

        column_names = [desc[0] for desc in c.description] if case_data else []
        conditions = {column_names[i]: bool(case_data[i]) for i in range(len(case_data)) if case_data[i] == 1}

        # Fetch children details from nok_details
        c.execute("""
            SELECT name, age, education FROM nok_details 
            WHERE army_number = ? AND LOWER(TRIM(relation)) IN ('son', 'daughter')
        """, (army_number,))
        children = [{"name": row[0], "age": row[1], "education": row[2]} for row in c.fetchall()]
        num_children = len(children)

        # Fetch HRA city for calculation
        c.execute("SELECT hra FROM basic_info WHERE army_number = ?", (army_number,))
        hra_data = c.fetchone()
        hra_city = hra_data[0] if hra_data else "Unknown"

        # Calculate Maintenance Allowance
        allowance, condition_percentages = calculate_maintenance(conditions, num_children, children, {"hra_city": hra_city})

        # Assign calculated allowances to each child
        for i, child in enumerate(children):
            if allowance["wife"] == "Eligible: 10%":
                child["allowance"] = assign_children_allowance(children, {"hra_city": hra_city})[i]
            else:
                child["allowance"] = "Eligible: 5.5%"

        conn.close()

        return jsonify({
            "success": True,
            "conditions": condition_percentages,
            "num_children": num_children,
            "wife_allowance": allowance.get("wife", "Eligible: 0%"),
            "children": children
        })

    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500


@app.route('/submit_wife_ma', methods=['POST'])
def submit_wife_ma():
    try:
        data = request.json
        army_number = data['army_number']
        selected_percentage = data['selected_percentage']

        print(f"\n[DEBUG] Received Wife MA Submission - Army Number: {army_number}, Selected Percentage: {selected_percentage}\n")

        conn = sqlite3.connect('database.db')
        c = conn.cursor()

        # **Create the wife_ma table if it doesn't exist**
        c.execute('''
            CREATE TABLE IF NOT EXISTS wife_ma (
                army_number TEXT PRIMARY KEY,
                maintenance_percentage TEXT
            )
        ''')

        print("[DEBUG] Checked/Created wife_ma table in database.")

        # **Insert or update wife's maintenance allowance**
        c.execute('''
            INSERT INTO wife_ma (army_number, maintenance_percentage)
            VALUES (?, ?) ON CONFLICT(army_number) DO UPDATE SET maintenance_percentage = ?
        ''', (army_number, selected_percentage, selected_percentage))

        conn.commit()
        conn.close()

        # **Store in session**
        session[f"wifeMA_{army_number}"] = selected_percentage
        print(f"[DEBUG] Wife MA stored in session: {session[f'wifeMA_{army_number}']}\n")

        return jsonify({'success': True, 'message': 'Wife MA submitted successfully!'})

    except Exception as e:
        print(f"[ERROR] Failed to submit Wife MA: {e}\n")
        return jsonify({'success': False, 'message': str(e)}), 500


@app.route('/submit_child_ma', methods=['POST'])
def submit_child_ma():
    try:
        data = request.json
        army_number = data['army_number']
        child_name = data['child_name']
        selected_percentage = data['selected_percentage']

        print(f"\n[DEBUG] Received Child MA Submission - Army Number: {army_number}, Child Name: {child_name}, Selected Percentage: {selected_percentage}\n")

        conn = sqlite3.connect('database.db')
        c = conn.cursor()

        # **Create the children_ma table if it doesn't exist**
        c.execute('''
            CREATE TABLE IF NOT EXISTS children_ma (
                army_number TEXT,
                child_name TEXT,
                maintenance_percentage TEXT,
                PRIMARY KEY (army_number, child_name)
            )
        ''')
        print("[DEBUG] Checked/Created children_ma table in database.")

        # **Insert or update child's maintenance allowance**
        c.execute('''
            INSERT INTO children_ma (army_number, child_name, maintenance_percentage)
            VALUES (?, ?, ?) ON CONFLICT(army_number, child_name) DO UPDATE SET maintenance_percentage = ?
        ''', (army_number, child_name, selected_percentage, selected_percentage))

        conn.commit()
        conn.close()

        # **Store in session**
        session[f"childMA_{army_number}_{child_name}"] = selected_percentage
        print(f"[DEBUG] Child MA stored in session: {session[f'childMA_{army_number}_{child_name}']}\n")

        return jsonify({'success': True, 'message': f'Maintenance Allowance for {child_name} submitted successfully!'})

    except Exception as e:
        print(f"[ERROR] Failed to submit Child MA: {e}\n")
        return jsonify({'success': False, 'message': str(e)}), 500



@app.route('/submit_all_children_ma', methods=['POST'])
def submit_all_children_ma():
    try:
        data = request.json
        army_number = data['army_number']
        children = data['children']

        print(f"\n[DEBUG] Received Final Submission for All Children - Army Number: {army_number}\n")
        print(f"[DEBUG] Children Data Received: {children}\n")

        conn = sqlite3.connect('database.db')
        c = conn.cursor()

        # **Create the children_ma table if it doesn't exist**
        c.execute('''
            CREATE TABLE IF NOT EXISTS children_ma (
                army_number TEXT,
                child_name TEXT,
                maintenance_percentage TEXT,
                PRIMARY KEY (army_number, child_name)
            )
        ''')
        print("[DEBUG] Checked/Created children_ma table in database.")

        # **Store or update maintenance allowance for each child**
        for child in children:
            child_name = child['child_name']
            selected_percentage = child['selected_percentage']

            print(f"[DEBUG] Processing Child: {child_name}, Selected Percentage: {selected_percentage}\n")

            c.execute('''
                INSERT INTO children_ma (army_number, child_name, maintenance_percentage)
                VALUES (?, ?, ?) ON CONFLICT(army_number, child_name) DO UPDATE SET maintenance_percentage = ?
            ''', (army_number, child_name, selected_percentage, selected_percentage))

            # **Store in session**
            session[f"childMA_{army_number}_{child_name}"] = selected_percentage
            print(f"[DEBUG] Child MA stored in session: {session[f'childMA_{army_number}_{child_name}']}\n")

        conn.commit()
        conn.close()

        print("[DEBUG] All Children MA submitted successfully!\n")
        return jsonify({'success': True, 'message': 'All children MA submitted successfully!'})

    except Exception as e:
        print(f"[ERROR] Failed to submit all children's MA: {e}\n")
        return jsonify({'success': False, 'message': str(e)}), 500
    
# THIS IS ROUTE FOR DOCUMENT GENERATION 


# @app.route("/get_information", methods=["POST"])
# def get_information():
#     print("in this route")
#     data = request.get_json()
#     army_number = data.get("army_number")

#     if not army_number:
#         return jsonify({"success": False, "message": "Army Number is required"}), 400

#     try:
#         # Database connection using context manager
#         with sqlite3.connect(db_path) as conn:
#             cursor = conn.cursor()

#             # Fetching basic information
#             cursor.execute("SELECT name, pcda_acct, unit, rank, date_of_marriage FROM basic_info WHERE army_number = ?", (army_number,))
#             basic_info_row = cursor.fetchone()
#             print(f"Fetched basic info: {basic_info_row}")
#             if not basic_info_row:
#                 return jsonify({"success": False, "message": "No record found for this Army Number."}), 404

#             name, pcda_acct, unit, rank, date_of_marriage = basic_info_row
            
#             # Fetching lady details
#             cursor.execute(""" 
#                 SELECT lady_address, lady_banker, lady_acct_number, lady_ifsc, date_of_affidavit 
#                 FROM lady_details WHERE army_number = ? 
#             """, (army_number,))
#             lady_details_row = cursor.fetchone()
#             print(f"Fetched lady details: {lady_details_row}")
            
#             lady_address, lady_banker, lady_acct_number, lady_ifsc, date_of_affidavit = lady_details_row if lady_details_row else ("N/A", "N/A", "N/A", "N/A")
            
#             # Fetching next of kin details
#             cursor.execute("""
#                 SELECT 
#                     CASE WHEN relation = 'Wife/Spouse' THEN name ELSE NULL END AS lady_name,
#                     CASE WHEN relation LIKE '%Daughter%' OR relation LIKE '%Son%' OR relation = 'Child' THEN name ELSE NULL END AS child_name,
#                     relation,
#                     date_of_birth
#                 FROM nok_details WHERE army_number = ?;
#             """, (army_number,))
#             nok_rows = cursor.fetchall()
#             print(f"Fetched NOK details: {nok_rows}")

#             lady_name = None
#             children = []
#             son_name=None
#             for row in nok_rows:
#                 if not lady_name and row[0]:  # If lady_name is not already assigned
#                     lady_name = row[0]  # Assign the wife/spouse's name
#                 child_name = row[1]  # The child's name
#                 relation = row[2] 
#                 date_of_birth = row[3]
#                 if relation and 'Daughter' in relation:
#                     daughter_name = child_name
#                     daughter_dob = date_of_birth
#                 elif relation and 'Son' in relation:
#                     son_name = child_name
#                     son_dob = date_of_birth

#                 # Collecting all child names
#                 if child_name:
#                     children.append(child_name)

#             children_str = ", ".join(children) if children else "None"
#             print(f"Children: {children_str}")

#             # Fetching legal case details
#             cursor.execute("SELECT date_of_complaint FROM legal_cases WHERE army_number = ?", (army_number,))
#             legal_case_row = cursor.fetchone()
#             date_of_complaint = legal_case_row[0] if legal_case_row else "N/A"
#             print(f"Fetched legal case info: {date_of_complaint}")
            
#             # Fetching document details
#             cursor.execute("SELECT date_reply_scn FROM documents WHERE army_number = ?", (army_number,))
#             document_row = cursor.fetchone()
#             date_reply_scn = document_row[0] if document_row else "N/A"
#             print(f"Fetched document details: {date_reply_scn}")

#             # Generate dynamic text for the document
#             dynamic_text = ""
#             if daughter_name:
#                 dynamic_text += f"In respect of the daughter Ms {daughter_name}, it will continue till she gets married or gains employment. "
#             if son_name:
#                 dynamic_text += f"In respect of the son Master {son_name}, it will continue till he attains 25 years of age or in the event of death of any of the concerned parties. "
#             dynamic_text += ("In case of the above conditions being met, recovery of the maintenance allowance will be discontinued by PAO (O)/ PAO (OR) concerned on receipt of intimation from the Commanding Officer/ Officer Commanding Unit or Establishment without the need of a separate order for discontinuation signed by the Competent Authority.")

#             # Call the function to summarize the case information
#             complaint_details = "Details of the complaint here..."
#             case_description = "Description of the case here..."
#             affidavit_filename = "C:/Users/Administrator/Downloads/random_affidavit.docx"  # Use the correct file path


# #THIS PATH FOR CASE FILE ADJUST AS PER YOUR FOLDER 
#             casefile = "C:/Users/Administrator/Desktop/latestcode/uploads/Petition_by_lady"

#             summary_result = summarize_case_info(complaint_details, case_description, affidavit_filename)

#             if not summary_result.get("success"):
#                 return jsonify({"success": False, "message": summary_result.get("message")}), 500

#             summarized_text = summary_result["summarized_text"]

#             # Check if template file exists
#             if not os.path.exists(template_path):
#                 return jsonify({"success": False, "message": "Template file not found."}), 500

#             # Load the Word template
#             doc = Document(template_path)

#             # Replace placeholders in the document
#             for para in doc.paragraphs:
#                 para.text = para.text.replace("{{name}}", name)
#                 para.text = para.text.replace("{{unit}}", unit)
#                 para.text = para.text.replace("{{rank}}", rank)
#                 para.text = para.text.replace("{{date_of_marriage}}", date_of_marriage)
#                 para.text = para.text.replace("{{army_number}}", army_number)
#                 para.text = para.text.replace("{{lady_name}}", lady_name if lady_name else "N/A")
#                 para.text = para.text.replace("{{child_name}}", children_str)
#                 para.text = para.text.replace("{{lady_address}}", lady_address)
#                 para.text = para.text.replace("{{lady_acct_number}}", lady_acct_number)
#                 para.text = para.text.replace("{{lady_ifsc}}", lady_ifsc)
#                 para.text = para.text.replace("{{date_of_affidavit}}", date_of_affidavit)
#                 para.text = para.text.replace("{{date_of_complaint}}", date_of_complaint)
#                 para.text = para.text.replace("{{date_reply_scn}}", date_reply_scn)
#                 para.text = para.text.replace("{{pcda_acct}}", pcda_acct)
#                 para.text = para.text.replace("{{lady_banker}}", lady_banker)
#                 para.text = para.text.replace("{{dynamic_text}}", dynamic_text)
#                 para.text = para.text.replace("{{summarized_text}}", summarized_text)
                

#             # Save to a BytesIO object
#             output = BytesIO()
#             doc.save(output)
#             output.seek(0)

#             return send_file(output, as_attachment=True, download_name=f"{army_number}_info.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

#     except Exception as e:
#         logging.error(f"Exception occurred: {e}")
#         return jsonify({"success": False, "message": str(e)}), 500
import os
import sqlite3
import logging
from io import BytesIO
from docx import Document
import ollama

# Set up logging
logging.basicConfig(level=logging.DEBUG)

# Initialize Flask app

# Define the database path and template path
db_path = 'database.db'
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        logging.error(f"Failed to extract text from DOCX file: {e}")
        return ""

# Function to summarize text using Ollama
# def summarize_text_with_ollama(text):
#     try:
#         if not text.strip():  # Check if the text is empty or consists of only whitespace
#             logging.error("The text to summarize is empty.")
#             return "Error: The provided text is empty and cannot be summarized."
        
#         # Request Ollama to summarize the given text (using 'llama3:latest')
#         response = ollama.chat(model="llama3:latest", messages=[{"role": "user", "content": f"Summarize the following: {text}"}])
#         return response['message']['content']
#     except Exception as e:
#         logging.error(f"Failed to summarize text with Ollama: {e}")
#         return f"Error summarizing text: {str(e)}"
def summarize_text_with_ollama(text):
    try:
        if not text.strip():
            logging.error("The text to summarize is empty.")
            return ""  # Return empty if no text is provided

        # Modified prompt to request a very lengthy and detailed summary
        prompt = (
            f"Provide a very lengthy, detailed, and comprehensive summarization of the following text. "
            f"Ensure all complaint details are fully included, without omitting anything important. "
            f"Expand on every point, and do not shorten or omit details. "
            f"Repeat key details where relevant to ensure the summary is exhaustive and informative: {text}"
        )

        # Request Ollama to summarize the given text
        response = ollama.chat(model="llama3:latest", messages=[{"role": "user", "content": prompt}])

        # Check if the response contains valid content, or if it's a message asking for more input
        summary_content = response.get('message', {}).get('content', "")
        
        # If the response contains an apology or similar message, return an empty string
        if "I apologize" in summary_content or not summary_content:
            logging.error("Ollama returned an apology or empty response.")
            return ""  # Return empty if the response is not valid

        # Log the response length to monitor if it's too long
        logging.debug(f"Received summary (length: {len(summary_content)} characters)")

        return summary_content
    except Exception as e:
        logging.error(f"Failed to summarize text with Ollama: {e}")
        return ""  # Return empty in case of an error


# Main function to summarize case info
def preprocess_text(text):
    """
    Cleans up and prepares text for summarization.
    """
    # Example steps: Remove unwanted whitespace, long irrelevant sections, etc.
    text = text.strip()  # Remove leading/trailing spaces
    text = text.replace("\n", " ")  # Replace newlines with spaces
    # Add other cleanup steps here, such as removing excess paragraphs or irrelevant content.
    return text

def summarize_case_info(complaint_details, case_description, affidavit_filename):
    try:
        logging.debug(f"Retrieved complaint_details: {complaint_details}")
        logging.debug(f"Retrieved case_description: {case_description}")

        # Step 1: Extract text from the affidavit file
        affidavit_text = ""
        if affidavit_filename:
            upload_affidavit_path = affidavit_filename
            logging.debug(f"Affidavit file path: {upload_affidavit_path}")

            # Check if file exists
            if not os.path.exists(upload_affidavit_path):
                logging.error(f"File not found: {upload_affidavit_path}")
                return {'success': False, 'message': 'Affidavit file not found.'}

            # Extract text from the affidavit
            affidavit_text = extract_text_from_docx(upload_affidavit_path)
            logging.debug(f"Extracted affidavit text: {affidavit_text[:100]}...")

        # Step 2: Preprocess text before summarization
        complaint_details = preprocess_text(complaint_details)
        case_description = preprocess_text(case_description)
        affidavit_text = preprocess_text(affidavit_text)

        # Step 3: Combine all text for summarization
        combined_text = f"Complaint Details: {complaint_details}\n\nCase Description: {case_description}\n\nAffidavit: {affidavit_text}"

        if not combined_text.strip():  # Ensure the combined text is not empty
            return {'success': False, 'message': 'The combined text is empty.'}

        # Step 4: Summarize each section separately
        complaint_summary = summarize_text_with_ollama(complaint_details)
        case_summary = summarize_text_with_ollama(case_description)
        affidavit_summary = summarize_text_with_ollama(affidavit_text)

        logging.debug(f"Summarized complaint: {complaint_summary}")
        logging.debug(f"Summarized case: {case_summary}")
        logging.debug(f"Summarized affidavit: {affidavit_summary}")

        # Step 5: Combine the summaries for final output
        final_summary = f"Complaint Summary: {complaint_summary}\n\nCase Summary: {case_summary}\n\nAffidavit Summary: {affidavit_summary}"

        # Return the summarized text
        return {'success': True, 'summarized_text': final_summary}

    except Exception as e:
        logging.error(f"Error in summarize_case_info: {e}")
        return {'success': False, 'message': str(e)}
@app.route("/get_information", methods=["POST"])
def get_information():
    print("in this route")
    data = request.get_json()
    army_number = data.get("army_number")

    if not army_number:
        return jsonify({"success": False, "message": "Army Number is required"}), 400

    try:
        with sqlite3.connect(db_path) as conn:
            cursor = conn.cursor()

            cursor.execute("SELECT name, pcda_acct, unit, rank, date_of_marriage FROM basic_info WHERE army_number = ?", (army_number,))
            basic_info_row = cursor.fetchone()
            if not basic_info_row:
                return jsonify({"success": False, "message": "No record found for this Army Number."}), 404

            name, pcda_acct, unit, rank, date_of_marriage = basic_info_row

            cursor.execute(""" 
                SELECT lady_address, lady_banker, lady_acct_number, lady_ifsc, date_of_affidavit 
                FROM lady_details WHERE army_number = ? 
            """, (army_number,))
            lady_details_row = cursor.fetchone()
            lady_address, lady_banker, lady_acct_number, lady_ifsc, date_of_affidavit = lady_details_row if lady_details_row else ("N/A", "N/A", "N/A", "N/A", "N/A")

            cursor.execute("""
                SELECT 
                    CASE WHEN relation = 'Wife/Spouse' THEN name ELSE NULL END AS lady_name,
                    CASE WHEN relation LIKE '%Daughter%' OR relation LIKE '%Son%' OR relation = 'Child' THEN name ELSE NULL END AS child_name,
                    relation,
                    date_of_birth
                FROM nok_details WHERE army_number = ?;
            """, (army_number,))
            nok_rows = cursor.fetchall()
            daughter_name=None
            lady_name = None
            children = []
            son_name=None
            for row in nok_rows:
                if not lady_name and row[0]:  # If lady_name is not already assigned
                    lady_name = row[0]  # Assign the wife/spouse's name
                child_name = row[1]  # The child's name
                relation = row[2] 
                if relation and 'Daughter' in relation:
                    daughter_name = child_name
                elif relation and 'Son' in relation:
                    son_name = child_name

                # Collecting all child names
                if child_name:
                    children.append(child_name)
            children_str = ", ".join(children) if children else "None"
#             print(f"Children: {children_str}")
               
            # lady_name = None
            # child_name = "N/A"
            # date_of_birth = "N/A"
            # for row in nok_rows:
            #     relation = row[1]
            #     if relation == "Wife/Spouse":
            #         lady_name = row[0]
            #     elif "Son" in relation or "Daughter" in relation or relation == "Child":
            #         child_name = row[0]
            #         print(child_name,"this is child")
            #         date_of_birth = row[2]

            cursor.execute("SELECT date_of_complaint FROM legal_cases WHERE army_number = ?", (army_number,))
            legal_case_row = cursor.fetchone()
            date_of_complaint = legal_case_row[0] if legal_case_row else "N/A"

            cursor.execute("SELECT date_reply_scn FROM documents WHERE army_number = ?", (army_number,))
            document_row = cursor.fetchone()
            date_reply_scn = document_row[0] if document_row else "N/A"

            complaint_details = "Details of the complaint here..."
            case_description = "Description of the case here..."
            affidavit_filename = "C:/Users/Administrator/Downloads/random_affidavit.docx"

            summary_result = summarize_case_info(complaint_details, case_description, affidavit_filename)

            if not summary_result.get("success"):
                return jsonify({"success": False, "message": summary_result.get("message")}), 500

            summarized_text = summary_result["summarized_text"]

            dynamic_text = ""
            if daughter_name:
                print(daughter_name,"this is duaghter name")
                dynamic_text += f"In respect of the daughter Ms {daughter_name}, it will continue till she gets married or gains employment. "
            if son_name:
                dynamic_text += f"In respect of the son Master {son_name}, it will continue till he attains 25 years of age or in the event of death of any of the concerned parties. "
            dynamic_text += "In case of the above conditions being met, recovery of the maintenance allowance will be discontinued by PAO (O)/ PAO (OR) concerned on receipt of intimation from the Commanding Officer/ Officer Commanding Unit or Establishment without the need of a separate order for discontinuation signed by the Competent Authority."

            final_text = f"""
GRANT OF MAINTENANCE ALLOWANCE TO
SMT {lady_name} WIFE OF RANK {rank} NAME {name} OF {unit} PCDAO (A/C) NO/ RECORDS OFFICE {pcda_acct}

________________________________________
IC {army_number} Rank {rank} Name {name} of {unit}
married Smt {lady_name} on {date_of_marriage}. The couple has son/daughter Miss/Master {child_name} born from the wedlock.
________________________________________
Smt {lady_name} vide her affidavit No {date_of_affidavit} affirmed that she has no source of income. A show cause notice was served to the officer to seek his reply on allegations. The Officer in his reply dt {date_reply_scn} stated …
AI GENERATION STARTS HERE
{summarized_text}
AI GENERATION END HERE
________________________________________
The couple have been counseled by various agencies, i.e., civil authorities and officers in the Army chain of Command. However, reconciliation could not take place. The Commanders in the chain of Command of the officer have examined the case and submitted their recommendations.
________________________________________
The case has been deliberated upon by the Competent Authority who has found adequate/inadequate/insufficient/compelling reasons for grant/non-grant of maintenance allowance to Smt {lady_name} and child/children.
________________________________________
Accordingly, in exercise of the powers conferred under Section 90(i)/ 91(i) of the Army Act, 1950 (Act 46 of 1950), read with Army Rule 193, as amended, GOC-in-C .............................. Command has accorded sanction for the deduction of ....................... % per month from the pay and allowances of Rk {rank} Name {name} and its payment to Smt/Shri/Miss {lady_name} and for the maintenance of Master {child_name} child/children commencing from the date of the application for maintenance, that is {date_of_complaint}.
________________________________________
in the following manner:
(a) …………………….. % for wife, Smt {lady_name}
(b) …………………….. % for daughter(s)/son(s) {child_name} 
________________________________________
The deduction of maintenance allowance will continue till the marriage of Smt {lady_name} with the officer/JCO/OR subsists, or till her death, or till death/discharge/retirement of the individual, for a maximum period of three years from the date of sanction of maintenance allowance, whichever is earlier. Smt {lady_name} is advised to file a civil suit for the grant of permanent maintenance allowance, a copy of which may be forwarded to HQ ................................within six months of receipt of this order. In case of delay in the grant of permanent maintenance allowance by Civil Courts, an extension of this maintenance allowance will be considered for a further period of two years, provided a civil suit for maintenance was filed within six months of this order.
________________________________________
{dynamic_text}

________________________________________
The amount, as and when deducted, shall be remitted at the expense of the officer/JCO/OR directly to the above-named person at the address given below or any other address she/he may intimate in due time.
Address & Bank Details of the lady
Name: Smt {lady_name}
Address: {lady_address}
Bank Acct No: {lady_acct_number}
Bank: {lady_banker}
IFSC: {lady_ifsc}
""".strip()

            return jsonify({"success": True, "final_text": final_text})

    except Exception as e:
        logging.error(f"Exception occurred: {e}")
        return jsonify({"success": False, "message": str(e)}), 500

import pandas as pd

@app.route('/match_historical_cases', methods=['POST'])
def match_historical_cases():
    try:
        data = request.json
        army_number = data.get('army_number')
        threshold = data.get('threshold', 50)  # Default threshold is 50%
        print(army_number, threshold)

        if not army_number:
            return jsonify({"success": False, "message": "Army number is required"}), 400

        conn = sqlite3.connect('database.db')
        c = conn.cursor()

        # Fetch required details from database
        c.execute("SELECT date_of_marriage FROM basic_info WHERE army_number = ?", (army_number,))
        date_of_marriage = c.fetchone()
        print(date_of_marriage,"this is date of marriage")
        if date_of_marriage:
            date_of_marriage = date_of_marriage[0]
            marriage_duration = (datetime(2024, 12, 31) - datetime.strptime(date_of_marriage, "%Y-%m-%d")).days // 365
        else:
            marriage_duration = 0

        c.execute("SELECT COUNT(*) FROM nok_details WHERE army_number = ? AND LOWER(TRIM(relation)) IN ('son', 'daughter')", (army_number,))
        num_children = c.fetchone()[0]

        c.execute("SELECT complaint_details FROM legal_cases WHERE army_number = ?", (army_number,))
        complaint_details = c.fetchone()
        print(complaint_details,"these are compplaint details")
        complaint_details = complaint_details[0] if complaint_details else "N/A"

        # Fetch boolean values from case_matrix
        c.execute("SELECT lady_matriculate, lady_intermediate, lady_graduate, lady_post_graduate, \
                    lady_earnings_salaried, lady_earnings_self_employed, lady_earnings_prior_work, \
                    liability_no, liability_personal_car, liability_home_loan FROM case_matrix \
                    WHERE army_number = ?", (army_number,))
        case_data = c.fetchone()
        print(case_data,"this is case data")
        education = ["Matriculate", "Intermediate", "Graduate", "Post Graduate"][case_data[:4].index(1)] if 1 in case_data[:4] else "N/A"
        employment = ["Salaried", "Self Employed", "Prior Work Experience"][case_data[4:7].index(1)] if 1 in case_data[4:7] else "N/A"
        financial_status = ["No Liability", "Personal/Car Loan", "Home Loan"][case_data[7:].index(1)] if 1 in case_data[7:] else "N/A"

        conn.close()

        # Load Historical Cases (Excel File)
        try:
            print("BEFORE FINDING THE SHEET")
            df = pd.read_excel('database.xlsx', sheet_name='Sheet1')
            df = df.fillna('nil')
            print("AFTER LOADING THE SHEET")
        except Exception as e:
            print("Error:", e)
        # Extract columns 4-9 for similarity
        cols_to_match = df.iloc[:, 3:9]
        cols_to_match['Duration of marriage'] = cols_to_match['Duration of marriage'].apply(lambda x: int(x.split()[0]) if 'years' in str(x) else 0)

        # Standardize numerical data
        scaler = StandardScaler()
        cols_to_match[['Duration of marriage', 'Number of children']] = scaler.fit_transform(cols_to_match[['Duration of marriage', 'Number of children']])
        # Encode categorical data
        text_cols = ['Allegation by the lady', 'Educational Qualification of the lady', 'lady\'s emp', 'fin liabilities of the indl.']
        label_encoder = LabelEncoder()
        cols_encoded = cols_to_match.copy()
        cols_encoded[text_cols] = cols_to_match[text_cols].apply(label_encoder.fit_transform)

        # Train Nearest Neighbors Model (Cosine Similarity)
        knn = NearestNeighbors(n_neighbors=5, metric='cosine')
        knn.fit(cols_encoded)
        # Prepare Input Data
        input_data = [[marriage_duration, num_children, complaint_details, education, employment, financial_status]]
        input_df = pd.DataFrame(input_data, columns=['Duration of marriage', 'Number of children', 'Allegation by the lady',
                                                     'Educational Qualification of the lady', 'lady\'s emp', 'fin liabilities of the indl.'])
        input_df[text_cols] = input_df[text_cols].apply(label_encoder.fit_transform)
        input_df[['Duration of marriage', 'Number of children']] = scaler.transform(input_df[['Duration of marriage', 'Number of children']])

        # Perform Similarity Search
        distances, indices = knn.kneighbors(input_df)
        similarities = (1 / (1 + distances)) * 100

        matched_cases = [{"case_number": df.iloc[idx, 1], "name": df.iloc[idx, 2], "ma_gtd": df.iloc[idx, 9]} for idx, sim in zip(indices[0], similarities[0]) if sim >= threshold]

        return jsonify({"success": True, "matched_cases": matched_cases})

    except Exception as e:
        return jsonify({"success": False, "message": str(e)}), 500



if __name__ == '__main__':
    app.run(debug=True)