# To achieve the requirement of summarizing the complaint_details, case_description, and upload_affidavit content using Ollama and adding debugging information, we will follow these steps:

# Extract the text from the complaint_details and case_description fields from the database directly.
# Extract the text from the uploaded affidavit file located in the 'uploads' folder.
# Use Ollama to summarize the extracted text.
# Return the summarized text along with debugging information.
# Steps and Code Implementation
# Here’s how you can implement the required functionality in your backend route:

# 1. Extract text from the database fields (complaint_details, case_description)
# This step will extract the relevant fields (complaint_details, case_description) directly from the basic_info table.

# 2. Extract text from the uploaded affidavit file
# For file handling, I’ll assume you’re dealing with DOCX files for the affidavit, as they are commonly used for documents. You will need to extract the text from the file using the python-docx library.

# 3. Use Ollama to summarize the content
# We’ll use Ollama’s API to summarize the content of the extracted text.

# 4. Debugging Information
# We'll print relevant debugging information at each step to ensure everything is being processed correctly.
# '
import os
import json
import sqlite3
from flask import request, jsonify, session
import ollama
from docx import Document

# Function to extract text from DOCX file (for affidavit)
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"[ERROR] Failed to extract text from DOCX file: {e}")
        return ""

# Function to summarize text using Ollama
def summarize_text_with_ollama(text):
    try:
        # Request Ollama to summarize the given text
        response = ollama.chat(messages=[{"role": "user", "content": f"Summarize the following: {text}"}])
        return response['text']
    except Exception as e:
        print(f"[ERROR] Failed to summarize text with Ollama: {e}")
        return "Error summarizing text."

@app.route('/summarize_case_info', methods=['POST'])
def summarize_case_info():
    try:
        # Step 1: Fetching the necessary data
        army_number = request.form['army-number']  # Assuming the army number is sent in the form data

        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        
        # Fetch complaint details and case description from the database
        c.execute('SELECT complaint_details, case_description FROM basic_info WHERE army_number = ?', (army_number,))
        result = c.fetchone()

        if not result:
            return jsonify({'success': False, 'message': 'No data found for the given Army Number.'})

        complaint_details, case_description = result

        print(f"[DEBUG] Retrieved complaint_details: {complaint_details}")
        print(f"[DEBUG] Retrieved case_description: {case_description}")

        # Step 2: Extract text from the affidavit file
        upload_affidavit_path = os.path.join('uploads', request.form['upload-affidavit'])  # Assuming the filename is passed in the form data
        print(f"[DEBUG] Affidavit file path: {upload_affidavit_path}")

        # Extract text from the uploaded affidavit (DOCX format assumed)
        affidavit_text = extract_text_from_docx(upload_affidavit_path)
        print(f"[DEBUG] Extracted affidavit text: {affidavit_text[:100]}...")  # Printing first 100 chars for debugging

        # Step 3: Combine all text for summarization
        combined_text = f"Complaint Details: {complaint_details}\n\nCase Description: {case_description}\n\nAffidavit: {affidavit_text}"

        # Step 4: Use Ollama to summarize the combined text
        summarized_text = summarize_text_with_ollama(combined_text)
        print(f"[DEBUG] Summarized text: {summarized_text}")

        # Return the summarized text
        return jsonify({'success': True, 'summarized_text': summarized_text})

    except Exception as e:
        print(f"[ERROR] Exception occurred: {e}")
        return jsonify({'success': False, 'message': str(e)})


# Accept all the required file uploads: outcome_counselling_file, effort_reconciliation_file, brief_case_file, reply_scn_file, etc.
# Handle the text extraction and OCR (where necessary) for each document type.
# Summarize the extracted text using Ollama.
# Save the file paths and metadata in the SQLite database.
# Return the summarized text for all files uploaded.



import os
import json
import sqlite3
from flask import request, jsonify
import ollama
import pytesseract
from PIL import Image
import fitz  # PyMuPDF for PDF handling
from docx import Document
from werkzeug.utils import secure_filename

# Helper function to extract text from a DOCX file
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"[ERROR] Failed to extract text from DOCX file: {e}")
        return ""

# Helper function to extract text from a PDF file
def extract_text_from_pdf(file_path):
    try:
        doc = fitz.open(file_path)  # Open the PDF file
        full_text = []
        for page in doc:
            full_text.append(page.get_text())  # Extract text from each page
        return "\n".join(full_text)
    except Exception as e:
        print(f"[ERROR] Failed to extract text from PDF file: {e}")
        return ""

# Helper function to extract text from a JPEG (or other image files) using OCR
def extract_text_from_image(file_path):
    try:
        # Open the image using Pillow
        img = Image.open(file_path)
        # Use Tesseract OCR to extract text from the image
        extracted_text = pytesseract.image_to_string(img)
        return extracted_text
    except Exception as e:
        print(f"[ERROR] Failed to extract text from image file: {e}")
        return ""

# Function to summarize the text using Ollama
def summarize_text_with_ollama(text):
    try:
        # Request Ollama to summarize the given text
        response = ollama.chat(messages=[{"role": "user", "content": f"Summarize the following: {text}"}])
        return response['text']
    except Exception as e:
        print(f"[ERROR] Failed to summarize text with Ollama: {e}")
        return "Error summarizing text."

@app.route('/submit_documents', methods=['POST'])
def submit_documents():
    try:
        print("\n[DEBUG] Received request to submit documents.\n")

        # File uploads - accepting multiple files
        outcome_counselling_file = request.files.get('outcome-counselling')
        effort_reconciliation_file = request.files.get('effort-reconciliation')
        brief_case_file = request.files.get('brief-case')
        reply_scn_file = request.files.get('reply-scn')

        # Manual date fields
        date_scn = request.form.get('date-scn', '')
        date_reply_scn = request.form.get('date-reply-scn', '')

        print(f"[DEBUG] Parsed Date Fields: SCN: {date_scn}, Reply SCN: {date_reply_scn}")

        # List to store summarized texts for each file
        summarized_texts = []

        # Handle the files and extract text
        files = [
            ('Outcome Counselling', outcome_counselling_file),
            ('Effort Reconciliation', effort_reconciliation_file),
            ('Brief Case', brief_case_file),
            ('Reply SCN', reply_scn_file)
        ]

        for file_label, uploaded_file in files:
            if uploaded_file:
                print(f"[DEBUG] Processing file: {file_label}")

                # Save the uploaded file
                file_path = save_file(uploaded_file)
                print(f"[DEBUG] File path stored for {file_label}: {file_path}")

                # Extract text from the uploaded file based on its type
                file_extension = os.path.splitext(file_path)[1].lower()
                extracted_text = ""
                
                if file_extension == '.docx':
                    extracted_text = extract_text_from_docx(file_path)
                elif file_extension == '.pdf':
                    extracted_text = extract_text_from_pdf(file_path)
                elif file_extension in ['.jpeg', '.jpg', '.png']:
                    extracted_text = extract_text_from_image(file_path)  # Using OCR for images
                else:
                    print(f"[ERROR] Unsupported file format for {file_label}: {file_extension}")
                    return jsonify({'success': False, 'message': f'Unsupported file format for {file_label}.'})

                print(f"[DEBUG] Extracted text from {file_label} (first 100 chars): {extracted_text[:100]}...")

                # Use Ollama to summarize the extracted text
                summarized_text = summarize_text_with_ollama(extracted_text)
                summarized_texts.append({file_label: summarized_text})

                print(f"[DEBUG] Summarized text for {file_label}: {summarized_text}")

                # Save the details to the database
                conn = sqlite3.connect('database.db')
                c = conn.cursor()

                # Ensure the table exists
                c.execute('''CREATE TABLE IF NOT EXISTS documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_label TEXT,
                    file_path TEXT,
                    date_scn TEXT,
                    date_reply_scn TEXT
                )''')

                inserted_data = {
                    "file_label": file_label,
                    "file_path": file_path,
                    "date_scn": date_scn,
                    "date_reply_scn": date_reply_scn
                }

                # Insert data into the database
                c.execute('''INSERT INTO documents (file_label, file_path, date_scn, date_reply_scn) 
                             VALUES (?, ?, ?, ?)''', tuple(inserted_data.values()))
                conn.commit()
                conn.close()

        # Return the summarized texts for all uploaded files
        return jsonify({'success': True, 'summarized_texts': summarized_texts, 'message': 'Documents uploaded and saved successfully!'})

    except Exception as e:
        print(f"[ERROR] Exception occurred: {e}")
        return jsonify({'success': False, 'message': str(e)})

# Helper function to save files
def save_file(uploaded_file):
    """ Save the uploaded file and return its file path. """
    if uploaded_file and uploaded_file.filename:
        file_path = os.path.join('uploads', secure_filename(uploaded_file.filename))
        uploaded_file.save(file_path)
        print(f"[DEBUG] File saved at: {file_path}")
        return file_path
    return None
